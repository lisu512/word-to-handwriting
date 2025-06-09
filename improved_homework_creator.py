#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os
import random
import argparse
import shutil
from pathlib import Path

def rand_select(option):
    """随机选择一个手写字体"""
    if not option:
        return None
    return random.choice(option)

def create_dir(dir_path):
    """创建目录"""
    if not os.path.exists(dir_path):
        os.makedirs(dir_path)
        print(f"创建目录: {dir_path}")

def check_dependencies():
    """检查依赖是否安装"""
    try:
        import docx
        import cv2
        print("✓ 依赖检查通过")
        return True
    except ImportError as e:
        print(f"❌ 缺少依赖: {e}")
        print("请安装: pip install python-docx opencv-python")
        return False

def create_word_document(text_lines, output_path):
    """创建Word文档，使用手写字体"""
    try:
        from docx import Document
        from docx.oxml.ns import qn
        from docx.shared import RGBColor, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("错误：需要安装 python-docx")
        return False

    # 字体文件路径
    script_dir = Path(__file__).parent
    font_dir = script_dir / "font"
    
    # 专门使用font目录下的4种真正手写字体
    handwriting_fonts = []
    font_paths = {}
    
    if font_dir.exists():
        font_files = list(font_dir.glob("*.ttf"))
        for font_file in font_files:
            font_name = font_file.stem
            handwriting_fonts.append(font_name)
            font_paths[font_name] = str(font_file)
    
    if not handwriting_fonts:
        print("错误：未找到手写字体文件")
        return False
    
    print(f"使用手写字体: {handwriting_fonts}")
    available_fonts = handwriting_fonts

    doc = Document()
    
    # 设置页面边距和行间距
    sections = doc.sections
    for section in sections:
        section.top_margin = Pt(100)     # 增大上边距
        section.bottom_margin = Pt(100)
        section.left_margin = Pt(150)    # 增大左边距
        section.right_margin = Pt(72)

    char_count = 0
    current_font = rand_select(available_fonts)  # 初始化字体
    lines_per_page = 25  # 每页约25行，确保多页分布
    current_line_count = 0
    
    for line in text_lines:
        line = line.strip()
        if not line:  # 空行作为段落分隔
            # 添加空段落
            para = doc.add_paragraph()
            current_line_count += 1
            continue
            
        # 检查是否需要分页
        if current_line_count >= lines_per_page:
            # 插入分页符
            doc.add_page_break()
            current_line_count = 0
        
        # 创建新段落
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # 设置段落格式
        from docx.shared import Inches
        paragraph_format = para.paragraph_format
        paragraph_format.line_spacing = 1.8  # 1.8倍行间距
        paragraph_format.space_after = Pt(6)  # 段后间距
        paragraph_format.first_line_indent = Inches(0.5)  # 首行缩进两字符
        
        # 写入段落内容
        for char in line:
            run = para.add_run(char)
            
            # 更智能的字体选择：相近字符使用相同字体，增加连贯性
            if char_count % random.randint(3, 8) == 0:  # 每3-8个字符换一种字体
                current_font = rand_select(available_fonts)
            
            run.font.name = current_font
            
            # 超大字体大小 (29-36磅)
            font_size = 32 + random.randint(-4, 4)
            run.font.size = Pt(font_size)
            
            # 更自然的颜色变化 (深黑到浅灰)
            color_base = random.choice([20, 35, 45, 60])  # 几种不同的基础颜色
            color_variation = random.randint(-15, 15)
            color_val = max(10, min(80, color_base + color_variation))
            run.font.color.rgb = RGBColor(color_val, color_val, color_val)
            
            # 随机字体效果
            if random.random() < 0.1:  # 10%概率加粗
                run.font.bold = True
            if random.random() < 0.05:  # 5%概率斜体
                run.font.italic = True
            
            # 设置中文字体
            run._element.rPr.rFonts.set(qn('w:eastAsia'), current_font)
            char_count += 1
        
        # 重置字体选择
        current_font = rand_select(available_fonts)  # 新段落重新选择字体
        current_line_count += 1
    
    doc.save(output_path)
    print(f"Word文档已保存: {output_path}")
    return True

def word_to_image(docx_path, output_dir):
    """将Word文档转换为图片"""
    try:
        # 尝试不同的转换方法
        if shutil.which("mutool"):
            return word_to_image_mutool(docx_path, output_dir)
        else:
            return word_to_image_alternative(docx_path, output_dir)
    except Exception as e:
        print(f"文档转换失败: {e}")
        return 0

def word_to_image_alternative(docx_path, output_dir):
    """替代的文档转换方法 - 支持多页"""
    try:
        # 尝试使用python-docx2txt + PIL
        print("使用替代方法转换文档...")
        from docx import Document
        from PIL import Image, ImageDraw, ImageFont
        
        doc = Document(docx_path)
        
        # 尝试加载多种字体
        script_dir = Path(__file__).parent
        font_dir = script_dir / "font"
        available_fonts = []
        font_objects = {}
        
        # 只使用指定的4个字体
        target_fonts = [ "MengMeiZiTi-1", "李国夫手写体", "白路俏丽手写体","千图纤墨体","pigtruman手写体"]
        
        if font_dir.exists():
            for font_file in font_dir.glob("*.ttf"):
                font_name = font_file.stem
                if font_name in target_fonts:
                    try:
                        font_obj = ImageFont.truetype(str(font_file), size=60)
                        available_fonts.append(font_name)
                        font_objects[font_name] = font_obj
                    except:
                        pass
        
        if not available_fonts:
            available_fonts = ['default']
            font_objects['default'] = ImageFont.load_default()
        
        print(f"加载了 {len(available_fonts)} 种字体: {available_fonts}")
        
        # 设置主字体为"李国夫手写体"，其他字体偶尔点缀
        main_font = "李国夫手写体" if "李国夫手写体" in available_fonts else available_fonts[0]
        current_font_name = main_font
        print(f"主字体: {main_font}")
        print(f"点缀字体: {[f for f in available_fonts if f != main_font]}")
        
                                 # 图片设置
        img_width, img_height = 1998, 2585
        line_height = 90  # 增大行高配合更大字体
        lines_per_page = 25  # 每页最多25行，适应自动换行后的行数
        page_count = 0
        
        current_page_lines = 0
        current_img = None
        current_draw = None
        y_position = 210  # 第一行红线位置，字体底部对齐
        
        def create_new_page():
            nonlocal current_img, current_draw, y_position, page_count, current_page_lines
            page_count += 1
            current_img = Image.new('RGB', (img_width, img_height), color='white')
            current_draw = ImageDraw.Draw(current_img)
            y_position = 120  # 字体底部对齐第一条红线位置
            current_page_lines = 0
        
        def save_current_page():
            if current_img:
                output_path = os.path.join(output_dir, f"text{page_count}.png")
                current_img.save(output_path)
                print(f"图片已保存: {output_path}")
        
        # 创建第一页
        create_new_page()
        
        # 直接从Word文档读取完整内容，保持段落结构
        lines = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                lines.append(paragraph.text.strip())
            else:
                # 保留空行作为段落分隔符
                lines.append("")
        
        def smart_wrap_text(text, max_chars=29):
            """严格控制29±2字符的智能换行"""
            if len(text) <= max_chars:
                return [text]
            
            lines = []
            remaining = text
            
            while len(remaining) > max_chars:
                # 默认在29字符处断开
                break_pos = max_chars
                
                # 只检查第30个字符（索引29）是否是标点符号
                if max_chars < len(remaining) and remaining[max_chars] in '，。！？；：""''（）【】《》、':
                    # 如果第30个字符是标点，在标点后断开
                    break_pos = max_chars + 1
                else:
                    # 否则严格在29字符处断开
                    break_pos = max_chars
                
                lines.append(remaining[:break_pos])
                remaining = remaining[break_pos:]
            
            if remaining:
                lines.append(remaining)
            
            return lines
        
        for i, line in enumerate(lines):
            line = line.strip()
            
            # 空行处理 - 段落间隔
            if not line:
                # 段落间留完整空行，确保与红线对齐
                y_position += line_height
                current_page_lines += 1
                continue
            
            # 严格控制29±2字符换行
            wrapped_lines = smart_wrap_text(line, max_chars=29)
            
            for j, wrapped_line in enumerate(wrapped_lines):
                # 检查是否需要新页面
                if current_page_lines >= lines_per_page:
                    save_current_page()
                    create_new_page()
                
                # 判断是否为段落首行（该段落的第一个包装行）
                is_paragraph_start = (j == 0)
                
                # 按字符绘制文本，每个字符随机选择字体
                x_offset = 150 if is_paragraph_start else 100  # 段落首行缩进
                current_x = x_offset
                
                for char in wrapped_line:
                    # 每个字符随机选择字体
                    if random.random() < 0.80:  # 40%概率使用其他字体
                        other_fonts = [f for f in available_fonts if f != main_font]
                        if other_fonts:
                            current_font_name = random.choice(other_fonts)
                        else:
                            current_font_name = main_font
                    else:
                        current_font_name = main_font
                    
                    # 字体基线调整 - 使用测试验证的偏移值
                    font_y_offsets = {
                        "MengMeiZiTi-1": 0,      # 基准字体
                        "千图纤墨体": -8,         # 向上调整8像素
                        "李国夫手写体": 0,        # 基准字体
                        "pigtruman手写体": 2,     # 向下调整2像素
                        "白路俏丽手写体": 1,      # 向下调整1像素
                    }
                    font_y_offset = font_y_offsets.get(current_font_name, 0)
                    
                    # 轻微随机偏移 - 减小范围保持真实感
                    x_random = random.randint(-1, 1)
                    y_random = random.randint(-1, 1)
                    
                    # 字体大小变化
                    size_variation = random.randint(-2, 2)
                    font_size = 60 + size_variation
                    varied_font = ImageFont.truetype(str(font_dir / f"{current_font_name}.ttf"), size=font_size)
                    
                    try:
                        # 绘制单个字符，应用字体基线调整
                        current_draw.text((current_x + x_random, y_position + y_random + font_y_offset), 
                                        char, font=varied_font, fill='black')
                        
                        # 计算字符宽度，移动到下一个字符位置
                        bbox = current_draw.textbbox((0, 0), char, font=varied_font)
                        char_width = bbox[2] - bbox[0]
                        current_x += char_width + random.randint(1, 3)  # 字符间距随机变化
                        
                    except:
                        # 备用方案
                        current_draw.text((current_x + x_random, y_position + y_random + font_y_offset), 
                                        char, font=font_objects[main_font], fill='black')
                        current_x += 30  # 固定宽度备用
                
                y_position += line_height
                current_page_lines += 1
        
        # 保存最后一页
        save_current_page()
        
        print(f"成功转换 {page_count} 页")
        return page_count
        
    except Exception as e:
        print(f"替代转换方法失败: {e}")
        return 0

def word_to_image_mutool(docx_path, output_dir):
    """使用mutool转换Word文档为图片"""
    try:
        from docx2pdf import convert
        from PyPDF2 import PdfFileReader
        
        # 先转换为PDF
        pdf_path = docx_path.replace('.docx', '.pdf')
        convert(docx_path, pdf_path)
        
        # 读取PDF页数
        reader = PdfFileReader(pdf_path)
        if reader.isEncrypted:
            reader.decrypt('')
        page_count = reader.getNumPages()
        
        # 转换为图片
        script_dir = Path(__file__).parent
        mutool_path = script_dir / "mutool.exe" if os.name == 'nt' else "mutool"
        
        for i in range(page_count):
            output_path = os.path.join(output_dir, f"text{i+1}.png")
            cmd = f'"{mutool_path}" draw -o "{output_path}" -w 1998 -h 2585 "{pdf_path}" {i+1}'
            os.system(cmd)
        
        return page_count
        
    except Exception as e:
        print(f"mutool转换失败: {e}")
        return word_to_image_alternative(docx_path, output_dir)

def create_lined_background(width, height):
    """创建带红线的作业本背景"""
    import cv2
    import numpy as np
    
    # 创建白色背景
    background = np.ones((height, width, 3), dtype=np.uint8) * 255
    
    # 红线颜色 (BGR格式，OpenCV使用BGR而不是RGB)
    red_color = (0, 0, 255)  # 纯红色
    line_thickness = 3
    
    # 绘制横线 - 调整间距和位置以配合字体对齐
    line_spacing = 90  # 调整行间距配合更大字体和行高
    start_y = 210  # 红线起始位置，从第一行文字位置开始
    for y in range(start_y, height - 100, line_spacing):
        cv2.line(background, (50, y), (width - 50, y), red_color, line_thickness)
    
    # 移除左边距线
    
    print(f"生成红线背景: {width}x{height}, 行间距: {line_spacing}")
    return background

def apply_background_and_effects(page_count, temp_dir, output_dir):
    """应用背景和自然效果"""
    try:
        import cv2
        import numpy as np
    except ImportError:
        print("错误：需要安装 opencv-python")
        return False
    
    script_dir = Path(__file__).parent
    background_path = script_dir / "background.JPG"
    
    # 直接生成带红线的信笺纸背景
    print("生成带红线的信笺纸背景")
    background = create_lined_background(1998, 2585)
    
    offset_x = 29
    offset_y = 100
    
    for num in range(page_count):
        word_img_path = os.path.join(temp_dir, f"text{num+1}.png")
        if not os.path.exists(word_img_path):
            continue
            
        word_img = cv2.imread(word_img_path)
        if word_img is None:
            continue
            
        # 不进行缩放，保持原始清晰度
        # word_img = cv2.resize(word_img, (0, 0), fx=1.1, fy=1.2)
        result = background.copy()
        
        print(f'处理第 {num+1}/{page_count} 页...')
        
        # 改进的文字叠加算法，保留红线
        for i in range(min(len(word_img), len(result) - offset_y)):
            for j in range(min(len(word_img[0]), len(result[0]) - offset_x)):
                try:
                    if i + offset_y < len(result) and j + offset_x < len(result[0]):
                        # 检查当前像素是否为文字(非白色)
                        word_pixel = word_img[i][j]
                        bg_pixel = result[i + offset_y][j + offset_x]
                        
                        # 如果文字像素不是白色，则叠加
                        if not (word_pixel[0] > 240 and word_pixel[1] > 240 and word_pixel[2] > 240):
                            # 保留背景的红线，只替换白色区域
                            if bg_pixel[2] < 100:  # 如果背景是红色(BGR中R值低)，保留红线
                                continue
                            else:
                                # 叠加文字到白色背景上
                                result[i + offset_y][j + offset_x] = word_pixel
                except:
                    pass
        
        # 保持信笺纸完全端正，不做任何变形或旋转
        # 移除所有旋转、缩放、平移效果
        
        # 去除噪声，保持图像清晰
        # 移除噪声添加，保证字体清晰
        
        # 保存结果
        output_path = os.path.join(output_dir, f"result_{num+1}.jpg")
        cv2.imwrite(output_path, result)
        print(f'✓ 第 {num+1} 页已保存: {output_path}')
    
    return True

def main():
    parser = argparse.ArgumentParser(description='改进版作业生成器 - 生成逼真的手写作业图片')
    parser.add_argument('input_file', nargs='?', default='input.txt', help='输入文本文件路径')
    parser.add_argument('-o', '--output', default='res', help='输出目录')
    parser.add_argument('-t', '--temp', default='temporary', help='临时文件目录')
    parser.add_argument('--text', help='直接输入文本内容')
    parser.add_argument('--auto', action='store_true', help='自动模式，跳过确认步骤')
    
    args = parser.parse_args()
    
    # 检查依赖
    if not check_dependencies():
        return
    
    script_dir = Path(__file__).parent
    temp_dir = script_dir / args.temp
    output_dir = script_dir / args.output
    
    # 创建必要目录
    create_dir(temp_dir)
    create_dir(output_dir)
    
    # 清理输出目录
    for file in output_dir.glob("*"):
        if file.is_file():
            file.unlink()
    
    # 获取文本内容和Word文档路径
    if args.text:
        text_lines = args.text.split('\n')
        # 生成Word文档
        docx_path = temp_dir / "text.docx"
        if not create_word_document(text_lines, docx_path):
            return
    else:
        input_path = script_dir / args.input_file
        if not input_path.exists():
            print(f"错误：输入文件不存在: {input_path}")
            return
        
        # 检查是否为docx文件
        if input_path.suffix.lower() == '.docx':
            print(f"直接使用Word文档: {input_path}")
            docx_path = input_path
            # 从Word文档读取内容用于显示信息
            try:
                from docx import Document
                doc = Document(input_path)
                text_lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
                print(f"Word文档包含 {len(text_lines)} 个段落")
            except:
                print("无法读取Word文档内容，但将继续处理")
                text_lines = ["Word文档内容"]
        else:
            # 处理txt文件
            with open(input_path, 'r', encoding='utf-8') as f:
                text_lines = f.readlines()
                        
            if not text_lines:
                print("错误：没有输入内容")
                return
            
            print(f"准备处理 {len(text_lines)} 行文本...")
            
            # 生成Word文档
            docx_path = temp_dir / "text.docx"
            if not create_word_document(text_lines, docx_path):
                return
    
    # 自动继续，无需用户确认
    print(f"\n生成的Word文档: {docx_path}")
    print("开始转换为手写图片...")
    
    # 转换为图片
    print("正在转换Word文档为图片...")
    page_count = word_to_image(docx_path, temp_dir)
    
    if page_count == 0:
        print("文档转换失败")
        return
    
    print(f"成功转换 {page_count} 页")
    
    # 应用背景和效果
    print("正在应用背景和自然效果...")
    if apply_background_and_effects(page_count, temp_dir, output_dir):
        print(f"\n✓ 完成！结果已保存到 {output_dir} 目录")
        print(f"生成了 {page_count} 张作业图片")
    else:
        print("应用效果失败")
    
    # 清理临时文件
    if not args.auto:
        clean = input("是否清理临时文件？(Y/n): ").strip().lower()
        if clean != 'n':
            for file in temp_dir.glob("*"):
                if file.is_file():
                    file.unlink()
            print("临时文件已清理")

if __name__ == '__main__':
    main() 